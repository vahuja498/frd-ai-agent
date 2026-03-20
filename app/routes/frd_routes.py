"""
app/routes/frd_routes.py
FastAPI routes — auto-saves every FRD to output/ as .md, .docx, .json
"""
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import FileResponse
from typing import Annotated
import json
import re
from datetime import date
from pathlib import Path

from app.models.schemas import FRDRequest, FRDResponse, HealthResponse
from app.services.ingestion import IngestionService
from app.services.vector_store import VectorStoreService
from app.services.retriever import RetrieverService
from app.services.llm_service import LLMService
from app.services.frd_generator import FRDGeneratorService
from app.services.validator import ValidatorService
from app.config import settings

router = APIRouter()

OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)

_vector_store: VectorStoreService = None
_llm_service: LLMService = None


def get_vector_store() -> VectorStoreService:
    return _vector_store


def get_llm_service() -> LLMService:
    return _llm_service


def init_services(vector_store: VectorStoreService, llm: LLMService) -> None:
    global _vector_store, _llm_service
    _vector_store = vector_store
    _llm_service = llm


def slugify(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9]+", "_", text).strip("_")


def save_frd_to_disk(project_name: str, frd_markdown: str, response: FRDResponse) -> dict:
    """Save FRD as .md, .docx, and .json in the output/ folder."""
    slug = slugify(project_name)
    today = date.today().isoformat()
    base = OUTPUT_DIR / f"{slug}_{today}"
    saved = {}

    # 1. Markdown
    md_path = Path(f"{base}.md")
    md_path.write_text(frd_markdown, encoding="utf-8")
    saved["markdown"] = str(md_path)

    # 2. JSON
    json_path = Path(f"{base}.json")
    json_path.write_text(
        json.dumps(response.model_dump(), indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    saved["json"] = str(json_path)

    # 3. DOCX
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        title_para = doc.add_paragraph(f"Functional Requirement Document\n{project_name}")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.size = Pt(20)
            run.bold = True

        doc.add_paragraph(f"Date: {today}  |  Version: 1.0  |  Status: Draft")
        doc.add_page_break()

        for line in frd_markdown.split("\n"):
            line = line.rstrip()
            if line.startswith("#### "):
                doc.add_heading(line[5:], level=4)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("|") or line.startswith("---"):
                doc.add_paragraph(line)
            elif line.strip():
                doc.add_paragraph(line)

        docx_path = Path(f"{base}.docx")
        doc.save(str(docx_path))
        saved["docx"] = str(docx_path)

    except Exception as e:
        saved["docx_error"] = str(e)

    print(f"\n[AutoSave] Files saved to: {OUTPUT_DIR.resolve()}")
    for k, v in saved.items():
        print(f"  {k:10} → {v}")

    return saved


# ── Routes ─────────────────────────────────────────────────────────────────────

@router.post(
    "/generate-frd",
    response_model=FRDResponse,
    summary="Generate and auto-save a complete FRD",
    tags=["FRD Generation"],
)
async def generate_frd(
    request: FRDRequest,
    vs: Annotated[VectorStoreService, Depends(get_vector_store)],
    llm: Annotated[LLMService, Depends(get_llm_service)],
) -> FRDResponse:
    if vs is None or llm is None:
        raise HTTPException(status_code=503, detail="Services not initialised")

    try:
        ingestion = IngestionService()
        combined_input = ingestion.combine_for_prompt(
            transcript=request.transcript,
            mom=request.mom,
            sow=request.sow,
        )

        retriever = RetrieverService(vs)
        retrieved = retriever.retrieve(
            transcript=request.transcript,
            mom=request.mom,
            sow=request.sow,
            top_k=settings.top_k_results,
        )
        rag_context = retriever.format_context(retrieved)

        generator = FRDGeneratorService(llm)
        frd_markdown = generator.generate(
            project_name=request.project_name,
            combined_input=combined_input,
            rag_context=rag_context,
        )

        validator = ValidatorService()
        validation_report, confidence_score = validator.validate(frd_markdown)

        response = FRDResponse(
            project_name=request.project_name,
            frd=frd_markdown,
            frd_structured=None,
            validation=validation_report,
            confidence_score=confidence_score,
            rag_sources_used=len(retrieved),
            model_used=llm.model_name,
        )

        # Auto-save to disk
        saved_files = save_frd_to_disk(request.project_name, frd_markdown, response)
        print(f"[AutoSave] Saved {len(saved_files)} file(s)")

        return response

    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"FRD generation failed: {str(exc)}")


@router.get(
    "/list-frds",
    summary="List all saved FRDs in the output folder",
    tags=["FRD Management"],
)
async def list_frds():
    """Returns all saved FRD files from the output/ folder."""
    files = sorted(OUTPUT_DIR.glob("*"))
    return {
        "output_folder": str(OUTPUT_DIR.resolve()),
        "files": [
            {
                "name": f.name,
                "type": f.suffix,
                "size_kb": round(f.stat().st_size / 1024, 1),
                "path": str(f),
            }
            for f in files if f.is_file()
        ],
        "total": len([f for f in files if f.is_file()]),
    }


@router.get(
    "/download-frd/{filename}",
    summary="Download a saved FRD file by filename",
    tags=["FRD Management"],
)
async def download_frd(filename: str):
    """Download any saved FRD file from the output/ folder."""
    file_path = OUTPUT_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"File not found: {filename}")
    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type="application/octet-stream",
    )


@router.get(
    "/health",
    response_model=HealthResponse,
    summary="Health check",
    tags=["Health"],
)
async def health_check(
    vs: Annotated[VectorStoreService, Depends(get_vector_store)],
) -> HealthResponse:
    return HealthResponse(
        status="ok",
        vectorstore_loaded=vs is not None,
        frd_count=vs.total_chunks if vs else 0,
        model_provider=settings.model_provider,
    )